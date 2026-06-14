from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOCX_PATH = OUT_DIR / "Debuggy_AI_Project_Report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(28, 32, 38)
MUTED = RGBColor(88, 96, 108)
LIGHT_FILL = "E8EEF5"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_inches) -> None:
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths_inches):
            cell = row.cells[index]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Debuggy AI")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(180, 12, 28)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Complete Project Report, Launch Guide, Features, and Working")
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    set_table_width(meta, [1.875, 4.625])
    rows = [
        ("Project type", "AI-powered software engineering assistant and code quality dashboard"),
        ("Frontend", "React, Vite, Tailwind CSS, lucide-react"),
        ("Backend", "FastAPI, Pydantic, Uvicorn"),
        ("Primary colors", "Professional dark dashboard with red and black visual system"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_fill(row.cells[0], LIGHT_FILL)
        row.cells[0].paragraphs[0].runs[0].bold = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        set_cell_fill(cell, LIGHT_FILL)
        cell.paragraphs[0].runs[0].bold = True
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = value
    set_table_width(table, widths)


def build() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "Debuggy AI is a modern developer dashboard that helps engineers review code, understand bugs, "
        "detect risky patterns, improve maintainability, estimate complexity, generate tests, and convert code between languages."
    )
    doc.add_paragraph(
        "The current project is not a plain HTML website anymore. It has been converted to a React frontend and a FastAPI backend."
    )

    doc.add_heading("2. Is It Really React and FastAPI?", level=1)
    add_table(
        doc,
        ["Layer", "Evidence in Repository", "Purpose"],
        [
            ["React frontend", "frontend/src/main.jsx", "Contains the React component tree, dashboard state, tool pages, and API calls."],
            ["Tailwind CSS", "frontend/src/styles.css and frontend/tailwind.config.js", "Defines utility-driven red/black styling, glass panels, responsive layout, and animations."],
            ["Vite", "frontend/package.json and frontend/vite.config.js", "Runs the React dev server and proxies API calls during development."],
            ["FastAPI backend", "backend/app.py", "Defines the API app, CORS middleware, Pydantic models, and analysis endpoints."],
            ["Python dependencies", "backend/requirements.txt", "Lists fastapi, uvicorn, pydantic, and python-dotenv."],
        ],
        [1.35, 2.4, 2.75],
    )

    doc.add_heading("3. How To Launch Locally", level=1)
    doc.add_paragraph("After cloning from GitHub, launch the backend and frontend in two separate terminals.")
    doc.add_heading("Backend terminal", level=2)
    doc.add_paragraph("cd backend", style="Intense Quote")
    doc.add_paragraph("pip install -r requirements.txt", style="Intense Quote")
    doc.add_paragraph("uvicorn app:app --reload --host 127.0.0.1 --port 8000", style="Intense Quote")
    doc.add_heading("Frontend terminal", level=2)
    doc.add_paragraph("cd frontend", style="Intense Quote")
    doc.add_paragraph("npm install", style="Intense Quote")
    doc.add_paragraph("npm run dev", style="Intense Quote")
    doc.add_paragraph("Open http://127.0.0.1:5173 in your browser.")

    doc.add_heading("4. GitHub and Deployment Notes", level=1)
    add_bullets(
        doc,
        [
            "Pushing to GitHub stores the source code; it does not automatically run the React or FastAPI servers.",
            "For local use, clone the repository and run the backend and frontend commands above.",
            "For public hosting, deploy the FastAPI backend on a service such as Render, Railway, Fly.io, or a VPS.",
            "Deploy the React frontend on Vercel, Netlify, or any static hosting provider that supports Vite builds.",
            "In production, route frontend /api requests to the FastAPI backend, or configure the frontend to call the backend URL directly.",
        ],
    )

    doc.add_heading("5. Application Architecture", level=1)
    add_numbered(
        doc,
        [
            "The user opens the React dashboard in the browser.",
            "The user writes or pastes code into the editor and selects a language.",
            "React sends JSON requests to the FastAPI API using /api routes in development.",
            "FastAPI analyzes the code with heuristic analysis functions and returns structured JSON.",
            "React renders errors, fixed code, explanations, scores, and advanced tool results dynamically.",
        ],
    )

    doc.add_heading("6. Main Dashboard Features", level=1)
    add_bullets(
        doc,
        [
            "Large code editor with line numbers and monospace formatting.",
            "Language selector for Python, Java, JavaScript, C, C++, C#, and Go.",
            "Analyze Code button with loading state.",
            "Errors Found section covering syntax errors, logic errors, runtime risks, bad practices, and security concerns.",
            "Suggested Fix section with fully corrected code and one-click copy.",
            "Error Explanation section explaining cause, reason, and prevention.",
            "Code Health Score from 0 to 100 with category and visual progress bar.",
            "Quick Summary showing total errors, security risks, complexity rating, and code quality rating.",
            "Download Analysis Report and Export Fixed Code actions.",
        ],
    )

    doc.add_heading("7. Advanced Analysis Tools", level=1)
    add_table(
        doc,
        ["Tool", "What It Analyzes", "What It Displays"],
        [
            ["Security Scanner", "SQL injection, command injection, hardcoded credentials, XSS, unsafe APIs, weak encryption, dangerous functions, authentication risks.", "Overall security score, risk level, findings, recommendations, and secure code."],
            ["Complexity Analyzer", "Loops, nested loops, branches, runtime growth, memory growth, and performance bottlenecks.", "Time complexity, space complexity, rating, bottlenecks, and optimization suggestions."],
            ["Test Case Generator", "Function intent and likely input classes.", "Unit tests, edge cases, invalid inputs, boundary conditions, stress tests, and copyable test block."],
            ["Code Smell Detector", "Duplicate code, long functions, deep nesting, poor names, dead code, debug artifacts, and maintainability issues.", "Severity, explanation, and improvement direction."],
            ["Code Converter", "Original code and selected source/target languages.", "Original code, converted code area, and copy button."],
        ],
        [1.45, 2.85, 2.2],
    )

    doc.add_heading("8. Backend API Endpoints", level=1)
    add_table(
        doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/", "GET", "Health check confirming the API is running."],
            ["/analyze", "POST", "Full dashboard analysis with errors, fixed code, explanations, score, summary, security, smells, and complexity."],
            ["/security", "POST", "Dedicated security analysis."],
            ["/complexity", "POST", "Dedicated complexity and performance analysis."],
            ["/tests", "POST", "Generated test case recommendations."],
            ["/smells", "POST", "Dedicated code smell and maintainability analysis."],
            ["/convert", "POST", "Code conversion response with original and converted code."],
        ],
        [1.5, 1.0, 4.0],
    )

    doc.add_heading("9. How The Analysis Works", level=1)
    doc.add_paragraph(
        "The backend currently uses local deterministic analysis heuristics. It scans code for common patterns such as unsafe dynamic execution, "
        "query concatenation, weak crypto names, hardcoded secrets, nested loops, assignment inside conditionals, bracket imbalance, and maintainability smells."
    )
    doc.add_paragraph(
        "This means the app works without requiring a paid AI key. Later, a real LLM provider can be connected behind the same FastAPI endpoints for deeper reasoning."
    )

    doc.add_heading("10. User Experience and Design", level=1)
    add_bullets(
        doc,
        [
            "Dark developer dashboard layout with red highlights and black surfaces.",
            "Glassmorphism cards, soft borders, glowing primary actions, and animated result cards.",
            "Responsive sidebar navigation for desktop and mobile.",
            "Professional engineering-tool feel instead of a simple debugging form.",
            "Copy buttons, report download, export fixed code, and interactive advanced pages.",
        ],
    )

    doc.add_heading("11. Recommended Future Improvements", level=1)
    add_bullets(
        doc,
        [
            "Add authentication and per-user analysis history.",
            "Connect an LLM provider for more accurate explanations and code conversion.",
            "Add file upload and GitHub repository scanning.",
            "Add real syntax highlighting with Monaco Editor or CodeMirror.",
            "Add project-level reports and PDF export directly from the browser.",
            "Add CI checks and automated backend/frontend test suites.",
        ],
    )

    doc.add_heading("12. Quick Troubleshooting", level=1)
    add_table(
        doc,
        ["Problem", "Likely Cause", "Fix"],
        [
            ["Frontend opens but analysis fails", "FastAPI backend is not running.", "Start backend on port 8000."],
            ["npm command fails", "Dependencies were not installed.", "Run npm install inside frontend."],
            ["uvicorn command not found", "Python dependencies missing or venv not active.", "Run pip install -r requirements.txt."],
            ["GitHub page does not run backend", "GitHub stores code but does not host FastAPI automatically.", "Deploy backend separately and route /api to it."],
        ],
        [1.8, 2.2, 2.5],
    )

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run("Debuggy AI - React + Tailwind frontend, FastAPI backend, professional AI code review dashboard.")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build()
    print(DOCX_PATH)
